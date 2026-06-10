# generate_docx.py
# Creates: Transfer_Pipeline_Documentation.docx
# Requires: pip install python-docx

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

TITLE = "Transfer Scanner → XML Processor → JSONL Exporter\nTechnical Documentation"

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold: run.bold = True
    if italic: run.italic = True
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

def add_numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style='List Number')

def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    # Apply a "Intense Quote" style for subtle shading if available
    try:
        p.style = 'Intense Quote'
    except:
        pass

def build_doc():
    doc = Document()

    # Base styles
    styles = doc.styles
    if 'Code' not in styles:
        s = styles.add_style('Code', WD_STYLE_TYPE.CHARACTER)
        s.font.name = 'Consolas'
        s.font.size = Pt(9)

    # Title
    title = doc.add_paragraph()
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, "This document describes the end-to-end pipeline for scanning files, processing XML/ZIP into JSONL, and exporting to a JSONL API. It also covers dockerization, logging, configuration, and operational guidance.")

    # Components
    add_heading(doc, "1. Components", 1)
    add_bullets(doc, [
        "xml_scanner.py — scans local/FTP, indexes files into SQLite, triggers processor and exporter.",
        "xml_processor.py — reads NEW files, validates XML/ZIP, produces JSONL batches into the good/ folder, emits jsonl_output_list.txt.",
        "jsonl_transfer.py — reads jsonl_output_list.txt, loads JSONL, sends to JSONL API via dblib._send_to_jsonl_server; deletes JSONLs after success.",
        "dblib.py — shared lib: config loader, SQLite ops, XML/ZIP processing, HTTP auth and upload, helpers."
    ])

    # Configuration
    add_heading(doc, "2. Configuration", 1)
    add_para(doc, "Primary configuration file: scanner_config.yaml (local vs Docker variants). Key fields:")
    add_bullets(doc, [
        "use_local / use_ftp — toggles local and FTP scanning.",
        "local.directory_to_scan — folder to watch (e.g., D:\\Transferfiles\\failed or /mnt/2scan).",
        "log_path — where logs are written.",
        "database_path — path to SQLite DB.",
        "good_folder — where JSONLs are stored.",
        "quarantine_folder — where problematic inputs are stored.",
        "jsonl_list_file — list of JSONL outputs for exporter.",
        "dbapi.host / dbapi.port — JSONL server hostname and port."
    ])
    add_para(doc, "Inside Docker, the loader normalizes paths and overrides dbapi.host to 'dbapi' so the exporter reaches the API container by DNS.")

    # Logging & Artifacts
    add_heading(doc, "3. Logging & Artifacts", 1)
    add_bullets(doc, [
        "Scanner log: logs/trans_scanner_YYYY_MM_DD_HH_MM_SS.log",
        "Processor log: logs/xml_processor_YYYYMMDD_HHMMSS.log",
        "Exporter log: logs/data_exporter_YYYYMMDD_HHMMSS.log",
        "SQLite DB: database_path (e.g., trans_index.db or /app/data/scanner.db)",
        "JSONLs: under good/ (e.g., good/zollner_YYYYMMDD_HHMMSS.jsonl)",
        "List for exporter: jsonl_output_list.txt",
        "Quarantine: under quarantine/ with details on failures."
    ])

    # Execution Flow
    add_heading(doc, "4. Execution Flow", 1)
    add_heading(doc, "4.1 xml_scanner.py", 2)
    add_bullets(doc, [
        "Loads config and ensures SQLite schema exists if DB is missing.",
        "Scans local and/or FTP based on flags.",
        "Inserts discovered files into xml_files/zip_files as status='new'.",
        "If new files found: runs xml_processor.py; if JSONLs produced: runs jsonl_transfer.py."
    ])
    add_heading(doc, "4.2 xml_processor.py", 2)
    add_bullets(doc, [
        "Loads 'new' files; ensures they still exist.",
        "Multithreaded processing: XMLs and ZIPs.",
        "Writes per-origin JSONL batches into good/.",
        "Writes jsonl_output_list.txt listing JSONLs for export."
    ])
    add_heading(doc, "4.3 jsonl_transfer.py", 2)
    add_bullets(doc, [
        "Reads jsonl_output_list.txt, loads each JSONL.",
        "Authenticates to /token, sends batches to /data (START ... batches ... END).",
        "On success, deletes each JSONL (and its sidecar manifest if you applied the post-upload marking patch).",
        "On failure, keeps JSONL(s) for retry."
    ])

    # Post-upload marking (if applied)
    add_heading(doc, "5. Post-Upload Marking (Recommended)", 1)
    add_para(doc, "To ensure crash-safety, mark rows as processed only after exporter upload succeeds:")
    add_bullets(doc, [
        "Processor no longer sets status='processed' for XMLs (nor inner XMLs from ZIPs).",
        "append_to_output_file writes a sidecar manifest (file.jsonl.manifest.jsonl) with {'name','path_incoming'} per record.",
        "Exporter reads the manifest after a successful upload and sets those rows to status='processed', then deletes JSONL + manifest.",
        "If the exporter/network fails, JSONL + manifest remain; DB rows stay not processed → safe to retry."
    ])

    # Local Run
    add_heading(doc, "6. Running Locally", 1)
    add_numbered(doc, [
        "Edit scanner_config.yaml: set local.directory_to_scan and dbapi.host.",
        "Run the scanner: python xml_scanner.py",
        "Check console/logs for: scanner added files → processor wrote JSONL → exporter sent and deleted JSONL."
    ])
    add_para(doc, "Quick SQL checks (SQLite):")
    add_code_block(doc, """\
-- pending summary:
SELECT status, COUNT(*) FROM xml_files GROUP BY status;

-- last few:
SELECT name, path_incoming, status
FROM xml_files
ORDER BY ROWID DESC LIMIT 20;""")

    # Dockerization
    add_heading(doc, "7. Dockerization", 1)
    add_para(doc, "Build the image from the repo root:")
    add_code_block(doc, "docker build -t scanner .")
    add_para(doc, "Run with volumes for persistence and input scanning:")
    add_code_block(doc, """\
docker run --rm -it \
  -v /path/on/host/data:/app/data \
  -v /path/on/host/scan:/mnt/2scan \
  -e JSONL_SERVER_HOST=dbapi \
  -e JSONL_SERVER_PORT=5444 \
  --name scanner \
  scanner""")
    add_para(doc, "Container paths (typical Docker config):")
    add_bullets(doc, [
        "/app/data/logs — logs",
        "/app/data/scanner.db — SQLite DB",
        "/app/data/good — JSONL outputs",
        "/app/data/quarantine — quarantined inputs",
        "/app/data/jsonl_output_list.txt — exporter list file",
        "/mnt/2scan — input scan folder"
    ])
    add_para(doc, "Inside Docker, dbapi.host is overridden to 'dbapi' by the loader when it detects container runtime.")

    # Optional docker-compose
    add_heading(doc, "8. Optional: docker-compose", 1)
    add_para(doc, "Example compose file to run scanner + API together:")
    add_code_block(doc, """\
version: "3.9"
services:
  api:
    image: your/jsonl-api
    container_name: dbapi
    ports:
      - "5444:5444"

  scanner:
    image: scanner
    container_name: scanner
    depends_on:
      - api
    environment:
      JSONL_SERVER_HOST: dbapi
      JSONL_SERVER_PORT: "5444"
    volumes:
      - /path/on/host/data:/app/data
      - /path/on/host/scan:/mnt/2scan
    restart: unless-stopped""")

    # Troubleshooting
    add_heading(doc, "9. Troubleshooting", 1)
    add_bullets(doc, [
        "“No new files” but there are files: confirm the loaded config path and values; ensure use_local:true and directory_to_scan points to the right folder.",
        "Exporter connection errors: verify dbapi.host/port, firewall; exporter falls back to HTTP if HTTPS fails.",
        "JSONLs not deleted: export failed; check exporter log. Files kept intentionally for retry.",
        "Quarantine growth: inspect quarantine logs for invalid ZIP, malformed XML, nested ZIP, lock timeouts."
    ])

    # Appendix: Key commands
    add_heading(doc, "Appendix: Quick Commands", 1)
    add_para(doc, "Local:")
    add_code_block(doc, "python xml_scanner.py")
    add_para(doc, "Tail logs (PowerShell):")
    add_code_block(doc, """\
Get-Content .\\logs\\xml_scanner_*.log -Wait
Get-Content .\\logs\\xml_processor_*.log -Wait
Get-Content .\\logs\\data_exporter_*.log -Wait""")
    add_para(doc, "List pending JSONLs:")
    add_code_block(doc, "dir good\\*.jsonl")

    # Save
    doc.save("Transfer_Pipeline_Documentation.docx")
    print("Saved: Transfer_Pipeline_Documentation.docx")

if __name__ == "__main__":
    build_doc()
