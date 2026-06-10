"""
Dash Replacer
Replaces all em dashes (—) with en dashes (–) in .txt, .csv, and .docx files.
Saves the result as a new file with _fixed suffix.
"""

import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

# ── docx support ──────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.oxml.ns import qn
    import copy
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


EM = "\u2014"   # —
EN = "\u2013"   # –


# ── replacement logic ─────────────────────────────────────────────────────

def fix_text_file(path: str) -> tuple[str, int]:
    """Replace em dashes in a plain text or CSV file."""
    with open(path, "r", encoding="utf-8") as fh:
        original = fh.read()
    count = original.count(EM)
    fixed = original.replace(EM, EN)
    out_path = _output_path(path)
    with open(out_path, "w", encoding="utf-8") as fh:
        fh.write(fixed)
    return out_path, count


def fix_docx_file(path: str) -> tuple[str, int]:
    """Replace em dashes in all text runs of a .docx file."""
    doc = Document(path)
    count = 0

    def fix_run(run):
        nonlocal count
        if EM in run.text:
            count += run.text.count(EM)
            run.text = run.text.replace(EM, EN)

    # body paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            fix_run(run)

    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        fix_run(run)

    # headers and footers
    for section in doc.sections:
        for hdr in (section.header, section.footer,
                    section.even_page_header, section.even_page_footer,
                    section.first_page_header, section.first_page_footer):
            if hdr is not None:
                for para in hdr.paragraphs:
                    for run in para.runs:
                        fix_run(run)

    out_path = _output_path(path)
    doc.save(out_path)
    return out_path, count


def _output_path(path: str) -> str:
    root, ext = os.path.splitext(path)
    return f"{root}_fixed{ext}"


# ── tkinter UI ────────────────────────────────────────────────────────────

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Dash Replacer  —  →  –")
        self.resizable(False, False)
        self._build_ui()
        self._center()

    def _center(self):
        self.update_idletasks()
        w, h = self.winfo_width(), self.winfo_height()
        x = (self.winfo_screenwidth()  - w) // 2
        y = (self.winfo_screenheight() - h) // 2
        self.geometry(f"{w}x{h}+{x}+{y}")

    def _build_ui(self):
        PAD = {"padx": 16, "pady": 10}

        # ── header ────────────────────────────────────────────────────
        hdr = tk.Frame(self, bg="#1F4E79")
        hdr.pack(fill="x")
        tk.Label(hdr, text="Dash Replacer", bg="#1F4E79", fg="white",
                 font=("Georgia", 16, "bold"), pady=12).pack()
        tk.Label(hdr, text="Replaces  \u2014  (em dash)  with  \u2013  (en dash)",
                 bg="#1F4E79", fg="#8FB4D4",
                 font=("Georgia", 10, "italic"), pady=4).pack()

        # ── file row ──────────────────────────────────────────────────
        frm = tk.Frame(self, pady=6)
        frm.pack(fill="x", **PAD)
        tk.Label(frm, text="File:", font=("Georgia", 11), width=5,
                 anchor="w").grid(row=0, column=0, sticky="w")

        self.path_var = tk.StringVar()
        entry = tk.Entry(frm, textvariable=self.path_var, width=48,
                         font=("Courier New", 10), state="readonly",
                         readonlybackground="#F5F8FA")
        entry.grid(row=0, column=1, padx=(4, 8))

        tk.Button(frm, text="Browse…", command=self._browse,
                  bg="#2E75B6", fg="white", font=("Georgia", 10),
                  relief="flat", cursor="hand2",
                  padx=10, pady=4).grid(row=0, column=2)

        # ── info label ────────────────────────────────────────────────
        self.info_var = tk.StringVar(value="Accepted formats: .txt  .csv  .docx")
        tk.Label(self, textvariable=self.info_var,
                 font=("Georgia", 10), fg="#595959",
                 justify="left").pack(anchor="w", padx=16)

        # ── progress bar (hidden until run) ───────────────────────────
        self.progress = ttk.Progressbar(self, mode="indeterminate",
                                        length=400)

        # ── run button ────────────────────────────────────────────────
        self.run_btn = tk.Button(self, text="Replace Dashes",
                                  command=self._run,
                                  bg="#C55A11", fg="white",
                                  font=("Georgia", 12, "bold"),
                                  relief="flat", cursor="hand2",
                                  padx=20, pady=8, state="disabled")
        self.run_btn.pack(pady=(8, 4))

        # ── result area ───────────────────────────────────────────────
        self.result_var = tk.StringVar()
        tk.Label(self, textvariable=self.result_var,
                 font=("Georgia", 10), fg="#1E6B3C",
                 wraplength=460, justify="left").pack(padx=16, pady=(4, 14))

    # ── callbacks ─────────────────────────────────────────────────────────

    def _browse(self):
        types = [
            ("Supported files", "*.txt *.csv *.docx"),
            ("Text files",      "*.txt"),
            ("CSV files",       "*.csv"),
            ("Word documents",  "*.docx"),
            ("All files",       "*.*"),
        ]
        path = filedialog.askopenfilename(title="Select a file", filetypes=types)
        if not path:
            return

        ext = os.path.splitext(path)[1].lower()
        if ext not in (".txt", ".csv", ".docx"):
            messagebox.showerror("Unsupported format",
                                 f"'{ext}' is not supported.\n"
                                 "Please choose a .txt, .csv, or .docx file.")
            return

        if ext == ".docx" and not DOCX_AVAILABLE:
            messagebox.showerror("Missing dependency",
                                 "python-docx is not installed.\n"
                                 "Run: pip install python-docx")
            return

        self.path_var.set(path)
        self.result_var.set("")
        self.info_var.set(f"Selected: {os.path.basename(path)}")
        self.run_btn.config(state="normal")

    def _run(self):
        path = self.path_var.get()
        if not path:
            return

        self.run_btn.config(state="disabled")
        self.result_var.set("Working…")
        self.progress.pack(pady=(0, 8))
        self.progress.start(12)
        self.update()

        try:
            ext = os.path.splitext(path)[1].lower()
            if ext == ".docx":
                out_path, count = fix_docx_file(path)
            else:
                out_path, count = fix_text_file(path)

            noun = "replacement" if count == 1 else "replacements"
            self.result_var.set(
                f"\u2713  {count} {noun} made.\n"
                f"Saved as: {os.path.basename(out_path)}"
            )
        except Exception as exc:
            self.result_var.set(f"\u2717  Error: {exc}")
            messagebox.showerror("Error", str(exc))
        finally:
            self.progress.stop()
            self.progress.pack_forget()
            self.run_btn.config(state="normal")


# ── entry point ───────────────────────────────────────────────────────────

if __name__ == "__main__":
    app = App()
    app.mainloop()